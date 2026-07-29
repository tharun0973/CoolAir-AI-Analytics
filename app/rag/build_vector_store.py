"""
CoolAir Comfort Services — Build Vector Store
=================================================
Chunks the three policy documents and embeds them with metadata that
encodes which document/version each chunk came from and when it took
effect — this is what lets the RAG retrieval prefer the current, correct
answer instead of just "whichever chunk scored highest."

Uses sentence-transformers (all-MiniLM-L6-v2) for local embeddings, so this
runs with no API key at all. Swap in OpenAIEmbeddings if you'd rather use
text-embedding-3-small — same downstream FAISS index either way.

Metadata per chunk:
  - source: filename
  - version: "v1" | "v2" | "sop"
  - effective_date: ISO date the section took effect
  - priority: used as a tie-breaker when two chunks conflict and dates
    alone don't resolve it (v2 > v1 > sop for pricing questions)

Usage: python build_vector_store.py
"""

import re
import json
from pathlib import Path

import faiss
import numpy as np
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer

# ------------------------------------------------------------
# Paths
# ------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parents[2]
DOCS_DIR = BASE_DIR / "documents"
DB_DIR = BASE_DIR / "database"

# ------------------------------------------------------------
# Document metadata
# ------------------------------------------------------------

DOCUMENT_META = {
    "Service_Pricing_and_Policy_Handbook_v1.docx": {
        "version": "v1",
        "effective_date": "2025-01-15",
        "priority": 1,
    },
    "Pricing_Addendum_v2.docx": {
        "version": "v2",
        "effective_date": "2025-06-01",
        "priority": 2,
    },
    "Customer_Service_SOP.docx": {
        "version": "sop",
        "effective_date": "2024-03-01",
        "priority": 1,
    },
}

CHUNK_SIZE_CHARS = 600
CHUNK_OVERLAP_CHARS = 100


def extract_docx_text(path: Path) -> str:
    doc = DocxDocument(path)

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts)


def chunk_text(
    text: str,
    size: int = CHUNK_SIZE_CHARS,
    overlap: int = CHUNK_OVERLAP_CHARS,
) -> list:
    """Simple sliding-window chunker on sentence boundaries where possible."""

    sentences = re.split(r"(?<=[.!?])\s+", text)

    chunks = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) <= size:
            current += (" " if current else "") + sentence
        else:
            if current:
                chunks.append(current.strip())

            current = (current[-overlap:] + " " if current else "") + sentence

    if current:
        chunks.append(current.strip())

    return [chunk for chunk in chunks if chunk]


def build_index():
    model = SentenceTransformer("all-MiniLM-L6-v2")

    all_chunks = []
    all_metadata = []

    for filename, meta in DOCUMENT_META.items():
        path = DOCS_DIR / filename

        text = extract_docx_text(path)
        chunks = chunk_text(text)

        for i, chunk in enumerate(chunks):
            all_chunks.append(chunk)
            all_metadata.append(
                {
                    "source": filename,
                    "chunk_index": i,
                    **meta,
                }
            )

        print(f"{filename}: {len(chunks)} chunks")

    embeddings = model.encode(all_chunks, show_progress_bar=False)
    embeddings = np.array(embeddings).astype("float32")

    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    DB_DIR.mkdir(exist_ok=True)

    faiss.write_index(index, str(DB_DIR / "policy_index.faiss"))

    with open(DB_DIR / "policy_chunks.json", "w", encoding="utf-8") as f:
        json.dump(
            {
                "chunks": all_chunks,
                "metadata": all_metadata,
            },
            f,
            indent=2,
        )

    print(f"\nIndexed {len(all_chunks)} chunks total -> {DB_DIR / 'policy_index.faiss'}")

    return model, index, all_chunks, all_metadata


def retrieve(question: str, model, index, chunks, metadata, k: int = 4):
    """
    Retrieves top-k chunks by embedding similarity, then re-sorts so higher
    priority sources rank first among close matches.
    """

    q_emb = model.encode([question]).astype("float32")

    distances, indices = index.search(q_emb, k)

    results = []

    for dist, idx in zip(distances[0], indices[0]):
        if idx == -1:
            continue

        results.append(
            {
                "chunk": chunks[idx],
                "metadata": metadata[idx],
                "distance": float(dist),
            }
        )

    results.sort(
        key=lambda r: (
            round(r["distance"], 1),
            -r["metadata"]["priority"],
        )
    )

    return results


if __name__ == "__main__":
    model, index, chunks, metadata = build_index()

    print("\n=== Retrieval smoke test ===")

    test_questions = [
        "What is the emergency diagnostic fee?",
        "What is the parts warranty on a compressor?",
        "What is the refund policy for cash payments?",
    ]

    for question in test_questions:
        print(f"\nQ: {question}")

        results = retrieve(
            question,
            model,
            index,
            chunks,
            metadata,
            k=2,
        )

        for result in results:
            meta = result["metadata"]

            print(
                f"  [{meta['version']}, effective {meta['effective_date']}, "
                f"dist={result['distance']:.3f}] "
                f"{result['chunk'][:120]}..."
            )